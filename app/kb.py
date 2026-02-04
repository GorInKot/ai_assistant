from __future__ import annotations

import math
import re
import subprocess
from collections import Counter, OrderedDict, defaultdict
from dataclasses import dataclass, field
from pathlib import Path
from urllib.parse import quote

try:
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover
    Document = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:  # pragma: no cover
    PdfReader = None


TOKEN_PATTERN = re.compile(r"[A-Za-zА-Яа-яЁё0-9]+")

SUPPORTED_EXTENSIONS = {".md", ".pdf", ".docx", ".doc"}

RU_STOPWORDS = {
    "и",
    "в",
    "во",
    "на",
    "по",
    "с",
    "со",
    "к",
    "ко",
    "от",
    "до",
    "из",
    "у",
    "о",
    "об",
    "не",
    "что",
    "как",
    "для",
    "или",
    "а",
    "но",
    "же",
    "это",
    "этот",
    "эта",
    "эти",
    "его",
    "ее",
    "их",
    "мы",
    "вы",
    "они",
    "я",
    "ты",
    "ли",
    "бы",
    "при",
    "над",
    "под",
    "про",
    "так",
    "если",
    "то",
    "есть",
    "какой",
    "какие",
    "какая",
    "какое",
    "где",
    "кто",
}

EN_STOPWORDS = {
    "the",
    "a",
    "an",
    "and",
    "or",
    "to",
    "in",
    "on",
    "at",
    "for",
    "of",
    "with",
    "by",
    "from",
    "is",
    "are",
    "be",
    "as",
    "that",
    "this",
    "it",
    "if",
    "then",
    "we",
    "you",
    "they",
    "what",
    "which",
    "where",
    "who",
}

RU_SUFFIXES = (
    "иями",
    "ями",
    "ами",
    "иями",
    "ов",
    "ев",
    "ого",
    "ему",
    "ыми",
    "ими",
    "ий",
    "ый",
    "ой",
    "ая",
    "ое",
    "ые",
    "ам",
    "ям",
    "ах",
    "ях",
    "ом",
    "ем",
    "ия",
    "ие",
    "иям",
    "иях",
    "а",
    "я",
    "ы",
    "и",
    "е",
    "о",
    "у",
    "ю",
)

EN_SUFFIXES = ("ing", "edly", "edly", "ed", "es", "s")

DEFINITION_HINTS = {
    "что такое",
    "что это",
    "для чего",
    "назначение",
    "what is",
    "purpose",
}

DOCUMENT_HINTS = {
    "документ",
    "форма",
    "шаблон",
    "бланк",
    "инструкция",
    "чеклист",
    "чек-лист",
    "document",
    "form",
    "template",
}

QUERY_NOISE_TOKENS = {
    "вопрос",
    "информац",
    "подскаж",
    "скаж",
    "расскаж",
    "мнен",
    "нужн",
    "хот",
    "пожалуйста",
    "please",
}

# Canonical semantic tags (syn:...) let us match paraphrases at scale
# without hard-coding logic for every question.
SYNONYM_GROUPS = {
    "participant_role": {
        "участник",
        "участники",
        "актор",
        "акторы",
        "роль",
        "роли",
        "исполнитель",
        "исполнители",
        "инициатор",
        "инициаторы",
        "заявитель",
        "заявители",
        "заказчик",
        "заказчики",
        "диспетчер",
        "диспетчеры",
        "контроллер",
        "контроллеры",
        "администратор",
        "администраторы",
        "специалист",
        "specialist",
        "participant",
        "participants",
        "actor",
        "actors",
        "role",
        "roles",
    },
    "status": {
        "статус",
        "статусы",
        "состояние",
        "state",
        "status",
        "statuses",
    },
    "application": {
        "заявка",
        "заявки",
        "обращение",
        "обращения",
        "request",
        "requests",
        "ticket",
        "tickets",
    },
    "step_checklist": {
        "шаг",
        "шаги",
        "порядок",
        "последовательность",
        "чеклист",
        "чек-лист",
        "step",
        "steps",
        "checklist",
    },
    "error_issue": {
        "ошибка",
        "ошибки",
        "проблема",
        "проблемы",
        "неисправность",
        "ошибоч",
        "error",
        "issue",
        "issues",
    },
}

ROLE_SPECIFIC_TERMS_RAW = {
    "заявитель",
    "заказчик",
    "диспетчер",
    "контроллер",
    "администратор",
    "специалист",
    "редактор",
    "исполнитель",
    "водитель",
}



def normalize_token_surface(token: str) -> str:
    value = token.lower().replace("ё", "е")
    if not value:
        return value
    if value.isdigit():
        return value

    if len(value) >= 5:
        for suffix in RU_SUFFIXES:
            if value.endswith(suffix) and len(value) - len(suffix) >= 3:
                value = value[: -len(suffix)]
                break

    if len(value) >= 5:
        for suffix in EN_SUFFIXES:
            if value.endswith(suffix) and len(value) - len(suffix) >= 3:
                value = value[: -len(suffix)]
                break

    return value



def _build_synonym_alias_map() -> dict[str, str]:
    alias_map: dict[str, str] = {}
    for alias, variants in SYNONYM_GROUPS.items():
        for variant in variants:
            normalized = normalize_token_surface(variant)
            if normalized:
                alias_map[normalized] = alias
    return alias_map


SYNONYM_ALIAS_BY_TOKEN = _build_synonym_alias_map()
DEFAULT_STOPWORDS = {normalize_token_surface(word) for word in (RU_STOPWORDS | EN_STOPWORDS)}
ROLE_SPECIFIC_TERMS = {normalize_token_surface(word) for word in ROLE_SPECIFIC_TERMS_RAW}


@dataclass
class TextUnit:
    text: str
    page: int | None = None
    section: str | None = None


@dataclass
class DocumentRecord:
    file_name: str
    relative_path: str
    absolute_path: Path
    process: str
    extension: str
    searchable: bool
    metadata_tokens: set[str] = field(default_factory=set)


@dataclass
class ChunkRecord:
    chunk_id: str
    relative_path: str
    text: str
    page: int | None
    section: str | None
    process: str
    token_counts: Counter[str]
    token_total: int
    section_tokens: set[str]
    path_tokens: set[str]
    role_mentions: int


@dataclass
class RetrievalResult:
    chunk: ChunkRecord
    score: float
    coverage: float


class KnowledgeBaseIndex:
    def __init__(self, kb_root: Path) -> None:
        self.kb_root = kb_root
        self.documents: dict[str, DocumentRecord] = {}
        self.chunks: list[ChunkRecord] = []
        self.token_idf: dict[str, float] = {}
        self.avg_chunk_len: float = 1.0
        self.process_tokens: dict[str, set[str]] = {}

    def build(self) -> None:
        self.documents = {}
        self.chunks = []
        self.token_idf = {}
        self.avg_chunk_len = 1.0
        self.process_tokens = {}

        if not self.kb_root.exists():
            return

        chunk_counter = 0
        process_tokens: dict[str, set[str]] = defaultdict(set)

        for file_path in sorted(self.kb_root.rglob("*")):
            if not file_path.is_file():
                continue
            if file_path.suffix.lower() not in SUPPORTED_EXTENSIONS:
                continue

            relative_path = file_path.relative_to(self.kb_root).as_posix()
            process = relative_path.split("/", 1)[0] if "/" in relative_path else "Общее"
            extension = file_path.suffix.lower()

            metadata_tokens = set(self._tokenize(f"{relative_path} {file_path.stem}"))
            process_tokens[process].update(self._tokenize(process))

            text_units: list[TextUnit] = []
            if extension == ".md":
                text_units = self._extract_md(file_path)
            elif extension == ".pdf":
                text_units = self._extract_pdf(file_path)
            elif extension == ".docx":
                text_units = self._extract_docx(file_path)
            elif extension == ".doc":
                text_units = self._extract_doc_best_effort(file_path)

            searchable = len(text_units) > 0
            self.documents[relative_path] = DocumentRecord(
                file_name=file_path.name,
                relative_path=relative_path,
                absolute_path=file_path,
                process=process,
                extension=extension,
                searchable=searchable,
                metadata_tokens=metadata_tokens,
            )

            if not searchable:
                continue

            for unit in text_units:
                section_tokens = set(self._tokenize(unit.section or ""))
                for chunk_text in self._chunk_text(unit.text):
                    indexed_text = f"{unit.section}\n{chunk_text}" if unit.section else chunk_text
                    plain_tokens = self._tokenize(indexed_text, include_alias=False)
                    tokens = Counter(self._tokenize(indexed_text))
                    if not tokens:
                        continue

                    role_mentions = len(set(plain_tokens) & ROLE_SPECIFIC_TERMS)

                    chunk_counter += 1
                    self.chunks.append(
                        ChunkRecord(
                            chunk_id=f"chunk_{chunk_counter}",
                            relative_path=relative_path,
                            text=indexed_text,
                            page=unit.page,
                            section=unit.section,
                            process=process,
                            token_counts=tokens,
                            token_total=sum(tokens.values()),
                            section_tokens=section_tokens,
                            path_tokens=metadata_tokens,
                            role_mentions=role_mentions,
                        )
                    )

        self.process_tokens = dict(process_tokens)
        self._recompute_index_stats()

    def retrieve(self, query: str, top_k: int = 20) -> list[RetrievalResult]:
        query_tokens = self._query_tokens(query)
        if not query_tokens:
            return []

        process_hint = self._detect_process_hint(query_tokens)
        strict_process_filter = self._should_strict_filter_by_process(query_tokens, process_hint)
        participant_query = "syn:participant_role" in query_tokens
        application_query = "syn:application" in query_tokens
        status_query = "syn:status" in query_tokens
        results: list[RetrievalResult] = []

        for chunk in self.chunks:
            if strict_process_filter and process_hint and chunk.process != process_hint:
                continue

            matched_tokens = query_tokens & set(chunk.token_counts.keys())
            if not matched_tokens:
                continue

            strong_matches = [token for token in matched_tokens if self.token_idf.get(token, 1.0) >= 1.45]
            if not strong_matches and len(matched_tokens) < 2:
                continue

            bm25_score = sum(self._bm25_term_score(chunk, token) for token in matched_tokens)
            coverage = len(matched_tokens) / max(len(query_tokens), 1)
            path_hits = len(query_tokens & chunk.path_tokens)
            section_hits = len(query_tokens & chunk.section_tokens)
            process_boost = 1.2 if process_hint and process_hint == chunk.process else 0.0
            markdown_boost = 0.1 if chunk.relative_path.endswith(".md") else 0.0
            role_boost = chunk.role_mentions * 1.05 if participant_query else 0.0
            semantic_combo_boost = 0.0
            if participant_query and application_query:
                if "syn:participant_role" in chunk.token_counts and "syn:application" in chunk.token_counts:
                    semantic_combo_boost += 1.4
            if status_query and chunk.section and "статус" in chunk.section.lower():
                semantic_combo_boost += 1.0

            score = (
                (bm25_score * 1.2)
                + (coverage * 3.0)
                + (path_hits * 0.55)
                + (section_hits * 0.9)
                + process_boost
                + markdown_boost
                + role_boost
                + semantic_combo_boost
            )
            if score <= 0:
                continue

            results.append(RetrievalResult(chunk=chunk, score=score, coverage=coverage))

        results.sort(key=lambda item: item.score, reverse=True)
        return results[:top_k]

    def detect_process_hint(self, query: str) -> str | None:
        return self._detect_process_hint(self._query_tokens(query))

    def classify_query_intent(self, query: str) -> str:
        query_lower = query.lower().strip()
        if any(hint in query_lower for hint in DEFINITION_HINTS):
            return "definition"
        if any(hint in query_lower for hint in DOCUMENT_HINTS):
            return "documents"
        return "procedure"

    def select_context_results(self, query: str, results: list[RetrievalResult], max_chunks: int = 5) -> list[RetrievalResult]:
        if not results:
            return []

        intent = self.classify_query_intent(query)
        query_tokens = self._query_tokens(query)
        candidate_results = results

        if intent == "definition":
            preferred = [
                item
                for item in results
                if item.chunk.relative_path.endswith("faq.md")
                or "glossary" in item.chunk.relative_path.lower()
                or (item.chunk.section and "назнач" in item.chunk.section.lower())
            ]
            if preferred:
                candidate_results = preferred + [item for item in results if item not in preferred]
            max_chunks = min(max_chunks, 3)

        selected: list[RetrievalResult] = []
        seen_chunks: set[str] = set()
        per_file: Counter[str] = Counter()
        per_file_limit = 1 if intent == "definition" else 2

        # Seed with chunks whose section heading semantically matches query terms.
        section_matched: list[tuple[int, float, RetrievalResult]] = []
        for item in candidate_results:
            overlap = len(query_tokens & item.chunk.section_tokens)
            if overlap > 0:
                section_matched.append((overlap, item.score, item))

        section_matched.sort(key=lambda x: (x[0], x[1]), reverse=True)
        for _, _, item in section_matched:
            if len(selected) >= max_chunks:
                break
            if item.chunk.chunk_id in seen_chunks:
                continue
            if per_file[item.chunk.relative_path] >= per_file_limit:
                continue
            selected.append(item)
            seen_chunks.add(item.chunk.chunk_id)
            per_file[item.chunk.relative_path] += 1

        for item in candidate_results:
            if len(selected) >= max_chunks:
                break
            if item.chunk.chunk_id in seen_chunks:
                continue
            if per_file[item.chunk.relative_path] >= per_file_limit:
                continue
            selected.append(item)
            seen_chunks.add(item.chunk.chunk_id)
            per_file[item.chunk.relative_path] += 1

        return selected if selected else results[:max_chunks]

    def is_context_strong(self, results: list[RetrievalResult]) -> bool:
        if not results:
            return False

        top = results[0]
        if top.score >= 2.8 and top.coverage >= 0.12:
            return True

        if len(results) >= 2:
            second = results[1]
            if (top.score + second.score) >= 5.0 and (top.coverage + second.coverage) >= 0.2:
                return True

        strong = [item for item in results if item.score >= 2.0]
        if len(strong) >= 3:
            return True

        return False

    def is_process_ambiguous(self, results: list[RetrievalResult], min_top_share: float = 0.62) -> bool:
        if len(results) < 2:
            return False

        by_process: dict[str, float] = defaultdict(float)
        for item in results:
            by_process[item.chunk.process] += max(item.score, 0.0)

        if len(by_process) <= 1:
            return False

        total = sum(by_process.values())
        if total <= 0:
            return False

        top_share = max(by_process.values()) / total
        return top_share < min_top_share

    def build_sources_from_results(self, results: list[RetrievalResult]) -> list[dict[str, object]]:
        grouped: OrderedDict[str, dict[str, object]] = OrderedDict()

        for result in results:
            chunk = result.chunk
            if chunk.relative_path not in grouped:
                grouped[chunk.relative_path] = {
                    "file_name": Path(chunk.relative_path).name,
                    "relative_path": chunk.relative_path,
                    "url": self._build_file_url(chunk.relative_path),
                    "download_url": self._build_file_url(chunk.relative_path, download=True),
                    "pages": set(),
                    "sections": set(),
                    "source_type": "context",
                }

            item = grouped[chunk.relative_path]
            if chunk.page is not None:
                item["pages"].add(chunk.page)
            if chunk.section:
                item["sections"].add(chunk.section)

        serialized: list[dict[str, object]] = []
        for item in grouped.values():
            serialized.append(
                {
                    "file_name": item["file_name"],
                    "relative_path": item["relative_path"],
                    "url": item["url"],
                    "download_url": item["download_url"],
                    "pages": sorted(item["pages"]),
                    "sections": sorted(item["sections"]),
                    "source_type": item["source_type"],
                }
            )

        return serialized

    def find_related_documents(self, query: str, limit: int = 3, exclude_paths: set[str] | None = None) -> list[dict[str, object]]:
        query_tokens = self._query_tokens(query)
        if not query_tokens:
            return []

        exclude_paths = exclude_paths or set()
        scored_docs: list[tuple[float, DocumentRecord]] = []

        for record in self.documents.values():
            if record.relative_path in exclude_paths:
                continue

            overlap_tokens = query_tokens & record.metadata_tokens
            if not overlap_tokens:
                continue

            score = sum(self.token_idf.get(token, 1.0) for token in overlap_tokens)
            if not record.searchable:
                score += 0.1

            scored_docs.append((score, record))

        scored_docs.sort(key=lambda item: item[0], reverse=True)

        related_sources: list[dict[str, object]] = []
        for _, record in scored_docs[:limit]:
            related_sources.append(
                {
                    "file_name": record.file_name,
                    "relative_path": record.relative_path,
                    "url": self._build_file_url(record.relative_path),
                    "download_url": self._build_file_url(record.relative_path, download=True),
                    "pages": [],
                    "sections": [],
                    "source_type": "related",
                }
            )

        return related_sources

    def list_documents(
        self,
        query: str | None = None,
        process: str | None = None,
        forms_only: bool = False,
        include_markdown: bool = False,
    ) -> list[dict[str, object]]:
        query_tokens = self._query_tokens(query or "")
        process_filter = (process or "").strip().lower()

        ranked: list[tuple[float, DocumentRecord]] = []
        for record in self.documents.values():
            if not include_markdown and record.extension == ".md":
                continue
            is_form = self._is_form_document(record)
            if forms_only and not is_form:
                continue
            if process_filter and process_filter not in record.process.lower():
                continue

            if query_tokens:
                overlap = query_tokens & record.metadata_tokens
                if not overlap:
                    continue
                score = sum(self.token_idf.get(token, 1.0) for token in overlap)
            else:
                score = 1.0

            if is_form:
                score += 0.1

            ranked.append((score, record))

        ranked.sort(key=lambda item: (-item[0], item[1].relative_path))

        response: list[dict[str, object]] = []
        for _, record in ranked:
            response.append(
                {
                    "file_name": record.file_name,
                    "relative_path": record.relative_path,
                    "url": self._build_file_url(record.relative_path),
                    "download_url": self._build_file_url(record.relative_path, download=True),
                    "process": record.process,
                    "extension": record.extension,
                    "searchable": record.searchable,
                    "is_form": self._is_form_document(record),
                }
            )
        return response

    def build_context_blocks(self, results: list[RetrievalResult], max_chars: int = 5500) -> list[dict[str, str | int | None]]:
        context: list[dict[str, str | int | None]] = []
        size = 0

        for result in results:
            chunk = result.chunk
            remaining = max_chars - size
            if remaining <= 0:
                break

            text = chunk.text.strip()
            if len(text) > remaining:
                text = text[:remaining]

            context.append(
                {
                    "text": text,
                    "relative_path": chunk.relative_path,
                    "page": chunk.page,
                    "section": chunk.section,
                }
            )
            size += len(text)

        return context

    def _extract_md(self, file_path: Path) -> list[TextUnit]:
        try:
            text = file_path.read_text(encoding="utf-8")
        except UnicodeDecodeError:
            text = file_path.read_text(encoding="cp1251", errors="ignore")

        return self._split_markdown_by_headings(text)

    def _extract_pdf(self, file_path: Path) -> list[TextUnit]:
        if PdfReader is None:
            return []

        try:
            reader = PdfReader(str(file_path))
        except Exception:
            return []

        units: list[TextUnit] = []
        for idx, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
            except Exception:
                page_text = ""

            page_text = page_text.strip()
            if not page_text:
                continue

            units.append(TextUnit(text=page_text, page=idx + 1))

        return units

    def _extract_docx(self, file_path: Path) -> list[TextUnit]:
        if Document is None:
            return []

        try:
            doc = Document(str(file_path))
        except Exception:
            return []

        units: list[TextUnit] = []
        section_name: str | None = None
        section_lines: list[str] = []

        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue

            style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
            is_heading = style_name.startswith("heading")

            if is_heading:
                if section_lines:
                    units.append(TextUnit(text="\n".join(section_lines), section=section_name))
                    section_lines = []
                section_name = text
                continue

            section_lines.append(text)

        if section_lines:
            units.append(TextUnit(text="\n".join(section_lines), section=section_name))

        if not units:
            full_text = "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())
            if full_text:
                units.append(TextUnit(text=full_text))

        return units

    def _extract_doc_best_effort(self, file_path: Path) -> list[TextUnit]:
        converters = [
            ["textutil", "-convert", "txt", "-stdout", str(file_path)],
            ["antiword", str(file_path)],
            ["catdoc", str(file_path)],
        ]

        for command in converters:
            try:
                result = subprocess.run(command, capture_output=True, text=True, timeout=8)
            except (FileNotFoundError, subprocess.SubprocessError):
                continue

            output = (result.stdout or "").strip()
            if result.returncode == 0 and len(output) >= 80:
                return [TextUnit(text=output)]

        for encoding in ("utf-8", "cp1251", "latin-1"):
            try:
                text = file_path.read_text(encoding=encoding, errors="ignore").strip()
            except Exception:
                continue
            if len(text) >= 120:
                return [TextUnit(text=text)]

        return []

    def _split_markdown_by_headings(self, text: str) -> list[TextUnit]:
        lines = text.splitlines()

        units: list[TextUnit] = []
        current_heading: str | None = None
        current_lines: list[str] = []

        for line in lines:
            stripped = line.strip()
            if stripped.startswith("#"):
                if current_lines:
                    units.append(TextUnit(text="\n".join(current_lines).strip(), section=current_heading))
                    current_lines = []
                current_heading = stripped.lstrip("#").strip() or current_heading
                continue

            current_lines.append(line)

        if current_lines:
            units.append(TextUnit(text="\n".join(current_lines).strip(), section=current_heading))

        if not units and text.strip():
            units.append(TextUnit(text=text.strip()))

        return [unit for unit in units if unit.text.strip()]

    def _chunk_text(self, text: str, max_len: int = 900, overlap: int = 120) -> list[str]:
        text = text.strip()
        if not text:
            return []
        if len(text) <= max_len:
            return [text]

        chunks: list[str] = []
        start = 0
        while start < len(text):
            end = min(start + max_len, len(text))
            piece = text[start:end].strip()
            if piece:
                chunks.append(piece)

            if end >= len(text):
                break

            start = max(0, end - overlap)

        return chunks

    def _tokenize(self, text: str, include_alias: bool = True) -> list[str]:
        tokens: list[str] = []
        for raw in TOKEN_PATTERN.findall(text):
            normalized = normalize_token_surface(raw)
            if not normalized or len(normalized) <= 1:
                continue
            if normalized in DEFAULT_STOPWORDS:
                continue

            tokens.append(normalized)
            if include_alias:
                alias = SYNONYM_ALIAS_BY_TOKEN.get(normalized)
                if alias:
                    tokens.append(f"syn:{alias}")

        return tokens

    def _query_tokens(self, query: str) -> set[str]:
        tokens = set(self._tokenize(query, include_alias=True))
        filtered = {token for token in tokens if token not in QUERY_NOISE_TOKENS}
        return filtered or tokens

    def _detect_process_hint(self, query_tokens: set[str]) -> str | None:
        best_process: str | None = None
        best_overlap = 0
        for process, tokens in self.process_tokens.items():
            overlap = len(query_tokens & tokens)
            if overlap > best_overlap:
                best_overlap = overlap
                best_process = process

        return best_process if best_overlap > 0 else None

    def _should_strict_filter_by_process(self, query_tokens: set[str], process_hint: str | None) -> bool:
        if not process_hint:
            return False
        process_terms = self.process_tokens.get(process_hint, set())
        explicit_terms = [
            token
            for token in (query_tokens & process_terms)
            if not token.startswith("syn:") and len(token) <= 5
        ]
        return bool(explicit_terms)

    def _bm25_term_score(self, chunk: ChunkRecord, token: str) -> float:
        tf = chunk.token_counts.get(token, 0)
        if tf <= 0:
            return 0.0

        idf = self.token_idf.get(token, 1.0)
        k1 = 1.2
        b = 0.75
        norm = 1.0 - b + b * (chunk.token_total / max(self.avg_chunk_len, 1.0))
        raw_score = idf * ((tf * (k1 + 1.0)) / (tf + (k1 * norm)))
        return min(raw_score, 3.5)

    def _recompute_index_stats(self) -> None:
        if not self.chunks:
            self.token_idf = {}
            self.avg_chunk_len = 1.0
            return

        doc_freq: Counter[str] = Counter()
        chunk_lens: list[int] = []

        for chunk in self.chunks:
            chunk_lens.append(chunk.token_total)
            for token in chunk.token_counts.keys():
                doc_freq[token] += 1

        total_chunks = len(self.chunks)
        self.token_idf = {
            token: math.log((total_chunks + 1.0) / (freq + 0.5)) + 1.0
            for token, freq in doc_freq.items()
        }
        self.avg_chunk_len = sum(chunk_lens) / max(len(chunk_lens), 1)

    def _build_file_url(self, relative_path: str, download: bool = False) -> str:
        encoded = quote(relative_path)
        if download:
            return f"/api/files/{encoded}?download=1"
        return f"/api/files/{encoded}"

    def _is_form_document(self, record: DocumentRecord) -> bool:
        name = record.file_name.lower()
        path = record.relative_path.lower()
        return any(hint in name or hint in path for hint in ("форма", "бланк", "шаблон", "form", "template"))
