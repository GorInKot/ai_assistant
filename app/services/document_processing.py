"""Обработка пользовательских документов (Этап 6 роадмапа).

Excel (`.xlsx`): извлечение данных (6.A.1), объединение нескольких файлов с
проверкой идентичности структуры (6.A.2), поиск различий между двумя файлами
(6.A.3).

Принципы:
- Stateless: файлы обрабатываются в памяти, на сервере не сохраняются.
  Это согласуется с локальным/безопасным характером продукта — корпоративные
  данные не оседают на диске.
- Без тяжёлых зависимостей: используем уже имеющийся `openpyxl`, без `pandas`.
- Лимиты на размер/строки/колонки, чтобы один большой файл не положил процесс.

Старый бинарный `.xls` вне скоупа первой итерации (openpyxl его не читает) —
для него отдаётся понятная ошибка с просьбой пересохранить в `.xlsx`.
"""

from __future__ import annotations

import difflib
import io
import os
import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field


# ---- Лимиты (значения по умолчанию; при нужде вынести в config) ----

# Лимиты размера разнесены по типам: openpyxl разворачивает .xlsx в объекты в
# памяти с коэффициентом ~10-40× от размера файла (узкое место RAM на одном
# процессе), тогда как текст .docx/.doc разворачивается скромнее. Поэтому Excel
# держим строже Word.
MAX_FILE_SIZE_XLSX = 15 * 1024 * 1024  # 15 МБ на Excel-файл
MAX_FILE_SIZE_WORD = 30 * 1024 * 1024  # 30 МБ на Word-файл (.docx/.doc)
MAX_ROWS = 50_000  # строк данных на лист (читаем не больше)
MAX_COLS = 256  # колонок на лист
MAX_FILES = 20  # файлов за одну операцию объединения/сравнения
PREVIEW_ROWS = 100  # строк в превью результата (полные данные — в скачиваемом файле)
MAX_DIFF_ITEMS = 1000  # ограничение размера отчёта о различиях
MAX_SUMMARY_CHARS = 30_000  # сколько символов текста отдаём в LLM на выжимку
DOC_CONVERT_TIMEOUT = 30  # сек на конвертацию старого .doc внешней утилитой
GRID_SCAN_ROWS = 1000  # сколько строк читаем «сеткой» при детекте формы (формы малы)
FORM_DENSITY_THRESHOLD = 0.30  # доля заполненных ячеек ниже которой лист считаем формой


class DocumentError(ValueError):
    """Ошибка обработки документа, безопасная для показа пользователю."""


def _check_size(content: bytes, limit: int) -> None:
    """Проверить размер файла против лимита (общий текст ошибки для xlsx/word)."""
    if len(content) > limit:
        raise DocumentError(
            f"Файл слишком большой ({len(content) // 1024} КБ). "
            f"Лимит — {limit // (1024 * 1024)} МБ."
        )


# ----------------------------- Модели данных -----------------------------


@dataclass
class SheetData:
    name: str
    headers: list[str]
    rows: list[list[str]]  # строки данных (без заголовка), уже обрезанные по лимиту
    total_rows: int  # всего строк данных в листе (до обрезки)
    truncated: bool  # True, если total_rows > len(rows)

    @property
    def n_cols(self) -> int:
        return len(self.headers)


@dataclass
class WorkbookData:
    filename: str
    sheets: list[SheetData]


@dataclass
class MergeResult:
    ok: bool
    headers: list[str]
    rows: list[list[str]]  # объединённые строки (полные)
    total_rows: int
    source_counts: list[tuple[str, int]]  # (filename, добавлено строк)
    mismatches: list[str] = field(default_factory=list)


@dataclass
class DiffResult:
    structure_changed: bool
    added_columns: list[str]
    removed_columns: list[str]
    reordered: bool
    key_column: str | None  # по какой колонке матчили строки (None → позиционно)
    added_rows: list[list[str]]
    removed_rows: list[list[str]]
    changed_cells: list[dict]  # {"key", "column", "old", "new"}
    truncated: bool  # отчёт обрезан по MAX_DIFF_ITEMS


# ------------------------------- Парсинг ---------------------------------


def _cell_to_str(value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, str):
        return value.strip()
    return str(value)


def parse_xlsx(
    content: bytes,
    filename: str,
    *,
    max_rows: int = MAX_ROWS,
    max_cols: int = MAX_COLS,
) -> WorkbookData:
    """Разобрать .xlsx в структурированное представление (листы/заголовки/строки)."""
    lower = (filename or "").lower()
    if lower.endswith(".xls") and not lower.endswith(".xlsx"):
        raise DocumentError(
            "Старый формат .xls не поддерживается. "
            "Пересохраните файл как .xlsx и загрузите снова."
        )
    if not lower.endswith(".xlsx"):
        raise DocumentError("Ожидается файл Excel в формате .xlsx")
    _check_size(content, MAX_FILE_SIZE_XLSX)

    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("openpyxl не установлен на сервере") from exc

    try:
        workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    except Exception as exc:  # noqa: BLE001 — openpyxl кидает разнородные ошибки
        raise DocumentError(f"Не удалось прочитать файл как .xlsx: {exc}") from exc

    sheets: list[SheetData] = []
    for ws in workbook.worksheets:
        # Материализуем строки (с запасом по лимиту), чтобы корректно определить
        # ширину таблицы и строку заголовков — частый случай: первая строка это
        # объединённый заголовок-«шапка» (напр. «Сотрудники»), а реальные имена
        # колонок идут ниже.
        raw_rows: list[list[str]] = []
        for raw in ws.iter_rows(values_only=True):
            raw_rows.append([_cell_to_str(v) for v in raw])
            if len(raw_rows) >= max_rows + 1:
                break

        # Истинная ширина = максимум по «последней непустой» ячейке среди строк.
        width = min(max((_last_nonempty_len(r) for r in raw_rows), default=0), max_cols)
        if width == 0:
            sheets.append(SheetData(name=ws.title, headers=[], rows=[], total_rows=0, truncated=False))
            continue

        header_idx = _detect_header_index(raw_rows, width)
        if header_idx is None:
            sheets.append(SheetData(name=ws.title, headers=[], rows=[], total_rows=0, truncated=False))
            continue

        headers = _fit_row(raw_rows[header_idx], width)
        # Пустым заголовкам даём подпись, чтобы колонка с данными не «исчезала».
        headers = [h if h else f"Колонка {i + 1}" for i, h in enumerate(headers)]

        data_rows: list[list[str]] = []
        total = 0
        for raw in raw_rows[header_idx + 1:]:
            cells = _fit_row(raw, width)
            if all(c == "" for c in cells):
                continue
            total += 1
            if len(data_rows) < max_rows:
                data_rows.append(cells)

        sheets.append(
            SheetData(
                name=ws.title,
                headers=headers,
                rows=data_rows,
                total_rows=total,
                truncated=total > len(data_rows),
            )
        )

    workbook.close()
    return WorkbookData(filename=filename, sheets=sheets)


def _last_nonempty_len(cells: list[str]) -> int:
    end = len(cells)
    while end > 0 and cells[end - 1] == "":
        end -= 1
    return end


def _fit_row(cells: list[str], width: int) -> list[str]:
    row = list(cells[:width])
    if len(row) < width:
        row += [""] * (width - len(row))
    return row


def _detect_header_index(rows: list[list[str]], width: int) -> int | None:
    """Найти строку заголовков, пропустив «титульные» строки сверху.

    Заголовок — первая строка, заполненная достаточно плотно для текущей ширины
    (для таблиц шире одной колонки — минимум 2 непустые ячейки; это отсекает
    одноячеечные титулы вроде объединённой «шапки»). Для одноколоночной таблицы
    берём первую непустую строку.
    """
    threshold = 1 if width == 1 else 2
    for idx, row in enumerate(rows):
        non_empty = sum(1 for c in row[:width] if c != "")
        if non_empty >= threshold:
            return idx
    return None


# --------------------------- Сравнение структуры -------------------------


def _norm_header(value: str) -> str:
    """Нормализация заголовка для сравнения: регистр и кратные пробелы игнорируются."""
    return " ".join(value.strip().lower().split())


def _structure_signature(headers: list[str]) -> list[str]:
    return [_norm_header(h) for h in headers]


# ----------------------------- 6.A.2 Объединение -------------------------


def merge_workbooks(workbooks: list[WorkbookData]) -> MergeResult:
    """Слить первые листы нескольких файлов в один.

    Предусловие: совпадение *набора* колонок по нормализованным именам (регистр
    и кратные пробелы игнорируются). **Порядок колонок может отличаться** — строки
    каждого файла автоматически переставляются к порядку первого файла (он задаёт
    итоговую структуру). При расхождении набора имён объединение не выполняется —
    возвращается, в каком файле каких колонок не хватает или какие лишние.
    """
    if len(workbooks) < 2:
        raise DocumentError("Для объединения нужно минимум два файла")

    base = workbooks[0].sheets[0] if workbooks[0].sheets else None
    if base is None or not base.headers:
        raise DocumentError(f"В файле «{workbooks[0].filename}» нет данных на первом листе")

    base_sig = _structure_signature(base.headers)
    base_set = set(base_sig)
    mismatches: list[str] = []
    # План переупорядочивания: для каждого файла — индексы его колонок в порядке base.
    plans: list[tuple[str, SheetData, list[int]]] = []

    for wb in workbooks:
        sheet = wb.sheets[0] if wb.sheets else None
        if sheet is None or not sheet.headers:
            mismatches.append(f"«{wb.filename}»: нет данных на первом листе")
            continue
        sig = _structure_signature(sheet.headers)
        if len(set(sig)) != len(sig):
            mismatches.append(
                f"«{wb.filename}»: повторяющиеся имена колонок — переупорядочивание неоднозначно"
            )
            continue
        sheet_set = set(sig)
        if sheet_set != base_set:
            missing = [h for h, s in zip(base.headers, base_sig) if s not in sheet_set]
            extra = [h for h, s in zip(sheet.headers, sig) if s not in base_set]
            parts = []
            if missing:
                parts.append(f"не хватает колонок: {', '.join(missing)}")
            if extra:
                parts.append(f"лишние колонки: {', '.join(extra)}")
            mismatches.append(f"«{wb.filename}»: {'; '.join(parts)}")
            continue
        pos = {s: i for i, s in enumerate(sig)}
        plans.append((wb.filename, sheet, [pos[s] for s in base_sig]))

    if mismatches:
        return MergeResult(
            ok=False,
            headers=base.headers,
            rows=[],
            total_rows=0,
            source_counts=[],
            mismatches=mismatches,
        )

    merged_rows: list[list[str]] = []
    source_counts: list[tuple[str, int]] = []
    for filename, sheet, order in plans:
        merged_rows.extend([row[i] for i in order] for row in sheet.rows)
        source_counts.append((filename, len(sheet.rows)))

    return MergeResult(
        ok=True,
        headers=base.headers,
        rows=merged_rows,
        total_rows=len(merged_rows),
        source_counts=source_counts,
    )


def build_xlsx(headers: list[str], rows: list[list[str]], *, sheet_name: str = "Merged") -> bytes:
    """Собрать .xlsx из заголовков и строк (для скачивания результата)."""
    from openpyxl import Workbook

    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name[:31] or "Sheet1"
    if headers:
        ws.append(headers)
    for row in rows:
        ws.append(row)
    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


# ------------------------------- 6.A.3 Различия --------------------------


def _pick_key_column(headers: list[str], a_rows: list[list[str]], b_rows: list[list[str]]) -> int | None:
    """Выбрать колонку-ключ для матчинга строк.

    Берём первую колонку, значения которой в ОБОИХ файлах уникальны и непусты.
    Если такой нет — None (тогда сравнение позиционное).
    """
    for col in range(len(headers)):
        a_vals = [r[col] for r in a_rows if col < len(r)]
        b_vals = [r[col] for r in b_rows if col < len(r)]
        if not a_vals or not b_vals:
            continue
        if any(v == "" for v in a_vals) or any(v == "" for v in b_vals):
            continue
        if len(set(a_vals)) == len(a_vals) and len(set(b_vals)) == len(b_vals):
            return col
    return None


def diff_workbooks(old: WorkbookData, new: WorkbookData) -> DiffResult:
    """Сравнить первые листы двух файлов: структура + строки/ячейки."""
    old_sheet = old.sheets[0] if old.sheets else None
    new_sheet = new.sheets[0] if new.sheets else None
    if old_sheet is None or new_sheet is None:
        raise DocumentError("Оба файла должны содержать хотя бы один лист с данными")

    old_sig = _structure_signature(old_sheet.headers)
    new_sig = _structure_signature(new_sheet.headers)
    old_set = set(old_sig)
    new_set = set(new_sig)

    added_columns = [h for h, s in zip(new_sheet.headers, new_sig) if s not in old_set]
    removed_columns = [h for h, s in zip(old_sheet.headers, old_sig) if s not in new_set]
    reordered = old_sig != new_sig and old_set == new_set
    structure_changed = bool(added_columns or removed_columns or reordered)

    truncated = False

    # Построчное сравнение делаем только по общим колонкам (по нормализованному имени).
    common = [s for s in old_sig if s in new_set]
    old_idx = {s: i for i, s in enumerate(old_sig)}
    new_idx = {s: i for i, s in enumerate(new_sig)}

    def project(row: list[str], idx_map: dict[str, int]) -> list[str]:
        out = []
        for s in common:
            i = idx_map[s]
            out.append(row[i] if i < len(row) else "")
        return out

    old_proj = [project(r, old_idx) for r in old_sheet.rows]
    new_proj = [project(r, new_idx) for r in new_sheet.rows]
    common_headers = [old_sheet.headers[old_idx[s]] for s in common]

    key_col = _pick_key_column(common_headers, old_proj, new_proj)
    added_rows: list[list[str]] = []
    removed_rows: list[list[str]] = []
    changed_cells: list[dict] = []

    def cap_reached() -> bool:
        nonlocal truncated
        if len(added_rows) + len(removed_rows) + len(changed_cells) >= MAX_DIFF_ITEMS:
            truncated = True
            return True
        return False

    if key_col is not None:
        old_by_key = {r[key_col]: r for r in old_proj}
        new_by_key = {r[key_col]: r for r in new_proj}
        for key, row in new_by_key.items():
            if key not in old_by_key and not cap_reached():
                added_rows.append(row)
        for key, row in old_by_key.items():
            if key not in new_by_key and not cap_reached():
                removed_rows.append(row)
        for key in old_by_key.keys() & new_by_key.keys():
            o, n = old_by_key[key], new_by_key[key]
            for col, header in enumerate(common_headers):
                if o[col] != n[col]:
                    if cap_reached():
                        break
                    changed_cells.append(
                        {"key": key, "column": header, "old": o[col], "new": n[col]}
                    )
    else:
        # Позиционное сравнение (нет надёжного ключа).
        for i in range(max(len(old_proj), len(new_proj))):
            o = old_proj[i] if i < len(old_proj) else None
            n = new_proj[i] if i < len(new_proj) else None
            if o is None and n is not None:
                if not cap_reached():
                    added_rows.append(n)
            elif n is None and o is not None:
                if not cap_reached():
                    removed_rows.append(o)
            elif o is not None and n is not None:
                for col, header in enumerate(common_headers):
                    if o[col] != n[col]:
                        if cap_reached():
                            break
                        changed_cells.append(
                            {"key": f"строка {i + 2}", "column": header, "old": o[col], "new": n[col]}
                        )

    return DiffResult(
        structure_changed=structure_changed,
        added_columns=added_columns,
        removed_columns=removed_columns,
        reordered=reordered,
        key_column=common_headers[key_col] if key_col is not None else None,
        added_rows=added_rows,
        removed_rows=removed_rows,
        changed_cells=changed_cells,
        truncated=truncated,
    )


# ----------------------- 6.A.3 (форма) Сравнение бланков -----------------
#
# Таблица-реестр (строка-заголовок + строки данных) сравнивается выше в
# `diff_workbooks`. Но заявки/бланки из Этапа 7 — это НЕ таблицы, а формы:
# пары «подпись — значение» с объединёнными ячейками и без строки заголовков.
# К ним табличный diff неприменим (колонки/строки бессмысленны). Поэтому для
# форм сравниваем ячейки по координатам, а каждое отличие подписываем ближайшей
# текстовой ячейкой слева/сверху (для наших бланков подпись лежит в колонке B).


@dataclass
class SheetGrid:
    """Лист как разрежённая сетка непустых ячеек (для сравнения форм)."""

    name: str
    cells: dict[tuple[int, int], str]  # (row, col) -> значение (только непустые)
    max_row: int
    max_col: int

    @property
    def density(self) -> float:
        area = self.max_row * self.max_col
        return len(self.cells) / area if area else 0.0


@dataclass
class FormDiffItem:
    label: str  # подпись поля (ближайшая текстовая ячейка слева/сверху)
    coord: str  # координата изменённой ячейки, напр. "F16"
    op: str  # "changed" | "added" | "removed"
    old: str | None
    new: str | None


@dataclass
class FormDiffResult:
    items: list[FormDiffItem]
    truncated: bool

    @property
    def summary(self) -> dict[str, int]:
        out = {"changed": 0, "added": 0, "removed": 0}
        for it in self.items:
            out[it.op] += 1
        return out


def parse_xlsx_grid(content: bytes, filename: str, *, max_rows: int = GRID_SCAN_ROWS) -> list[SheetGrid]:
    """Прочитать листы как сетку непустых ячеек с координатами (для форм).

    Читаем не больше `max_rows` строк: формы малы, а у больших таблиц этого
    хватит, чтобы оценить плотность и не зачитывать весь реестр впустую.
    """
    try:
        from openpyxl import load_workbook
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("openpyxl не установлен на сервере") from exc

    try:
        workbook = load_workbook(io.BytesIO(content), data_only=True, read_only=True)
    except Exception as exc:  # noqa: BLE001
        raise DocumentError(f"Не удалось прочитать файл как .xlsx: {exc}") from exc

    grids: list[SheetGrid] = []
    for ws in workbook.worksheets:
        cells: dict[tuple[int, int], str] = {}
        max_r = max_c = 0
        for r_idx, row in enumerate(ws.iter_rows()):
            if r_idx >= max_rows:
                break
            for cell in row:
                v = _cell_to_str(cell.value)
                if v != "":
                    cells[(cell.row, cell.column)] = v
                    max_r = max(max_r, cell.row)
                    max_c = max(max_c, cell.column)
        grids.append(SheetGrid(name=ws.title, cells=cells, max_row=max_r, max_col=max_c))

    workbook.close()
    return grids


def is_form_like(grid: SheetGrid) -> bool:
    """Эвристика «форма vs таблица»: форма — разрежённый лист.

    У таблицы-реестра заполнено большинство ячеек прямоугольного блока, у формы
    значения раскиданы по странице среди подписей и объединённых ячеек — плотность
    низкая. Порог `FORM_DENSITY_THRESHOLD` уверенно разделяет наши заявки (~0.02-0.05)
    и реестры сотрудников (плотность близка к 1).
    """
    if grid.max_row == 0 or grid.max_col == 0:
        return False
    return grid.density < FORM_DENSITY_THRESHOLD


def _nearest_label(grid: SheetGrid, row: int, col: int) -> str:
    """Подпись для ячейки: ближайшая непустая ячейка слева в строке, иначе сверху."""
    left = [(c, v) for (r, c), v in grid.cells.items() if r == row and c < col]
    if left:
        return max(left, key=lambda t: t[0])[1]
    above = [(r, v) for (r, c), v in grid.cells.items() if c == col and r < row]
    if above:
        return max(above, key=lambda t: t[0])[1]
    return ""


def _coord(row: int, col: int) -> str:
    from openpyxl.utils import get_column_letter

    return f"{get_column_letter(col)}{row}"


def diff_forms(old: SheetGrid, new: SheetGrid) -> FormDiffResult:
    """Сравнить две формы по ячейкам с привязкой к подписям (6.A.3 для бланков)."""
    items: list[FormDiffItem] = []
    truncated = False
    for key in sorted(set(old.cells) | set(new.cells)):
        o = old.cells.get(key, "")
        n = new.cells.get(key, "")
        if o == n:
            continue
        if len(items) >= MAX_DIFF_ITEMS:
            truncated = True
            break
        row, col = key
        label = _nearest_label(old if o else new, row, col)
        op = "changed" if (o and n) else ("added" if n else "removed")
        items.append(
            FormDiffItem(label=label, coord=_coord(row, col), op=op, old=o or None, new=n or None)
        )
    return FormDiffResult(items=items, truncated=truncated)


# ============================ Word (.docx) ===============================


@dataclass
class DocxParagraph:
    text: str
    style: str  # имя стиля (Normal, Heading 1, …)
    is_heading: bool


@dataclass
class DocxTable:
    rows: list[list[str]]


@dataclass
class DocxData:
    filename: str
    paragraphs: list[DocxParagraph]
    tables: list[DocxTable]

    @property
    def full_text(self) -> str:
        return "\n".join(p.text for p in self.paragraphs if p.text)


def _doc_to_text(content: bytes) -> str:
    """Извлечь текст из старого бинарного .doc (Word 97-2003).

    python-docx читает только .docx, поэтому .doc конвертируем внешней утилитой.
    Бэкенды пробуются по доступности (`shutil.which`):
      - `antiword` — Linux/прод (ставится в Docker), быстрый, текст без таблиц;
      - `textutil` — macOS/dev (встроен в систему), `.doc` → txt.
    Таблицы и форматирование не извлекаются — осознанное упрощение: все наши
    Word-операции текстовые (сравнение, выжимка, извлечение), а .doc — легаси.
    Конвертация требует временного файла на диске (утилиты не читают stdin для
    бинарных форматов); каталог удаляется сразу после — stateless сохраняется.
    """
    with tempfile.TemporaryDirectory(prefix="doc_conv_") as tmp:
        src = os.path.join(tmp, "in.doc")
        with open(src, "wb") as f:
            f.write(content)

        antiword = shutil.which("antiword")
        if antiword:
            try:
                res = subprocess.run(
                    [antiword, src], capture_output=True, timeout=DOC_CONVERT_TIMEOUT
                )
            except subprocess.TimeoutExpired as exc:
                raise DocumentError("Конвертация .doc заняла слишком долго") from exc
            if res.returncode == 0 and res.stdout.strip():
                return res.stdout.decode("utf-8", errors="replace")
            # antiword не справился (например, это не настоящий .doc) — пробуем дальше

        textutil = shutil.which("textutil")
        if textutil:
            dst = os.path.join(tmp, "out.txt")
            try:
                res = subprocess.run(
                    [textutil, "-convert", "txt", "-encoding", "UTF-8", "-output", dst, src],
                    capture_output=True,
                    timeout=DOC_CONVERT_TIMEOUT,
                )
            except subprocess.TimeoutExpired as exc:
                raise DocumentError("Конвертация .doc заняла слишком долго") from exc
            if res.returncode == 0 and os.path.exists(dst):
                with open(dst, "rb") as f:
                    return f.read().decode("utf-8", errors="replace")

        raise DocumentError(
            "Не удалось обработать .doc: на сервере нет конвертера (antiword/textutil). "
            "Пересохраните файл как .docx и загрузите снова."
        )


def _parse_doc_legacy(content: bytes, filename: str) -> DocxData:
    """Разобрать старый .doc: только текст (по строкам), без таблиц.

    Заголовки из plain-text надёжно не восстановить, поэтому `is_heading=False`
    у всех абзацев — извлечение покажет текст, выжимка/сравнение работают как
    обычно (они и так текстовые).
    """
    text = _doc_to_text(content)
    paragraphs = [
        DocxParagraph(text=line.strip(), style="Normal", is_heading=False)
        for line in text.splitlines()
        if line.strip()
    ]
    if not paragraphs:
        raise DocumentError("В файле .doc не найдено текста (возможно, повреждён или пуст)")
    return DocxData(filename=filename, paragraphs=paragraphs, tables=[])


def parse_docx(content: bytes, filename: str) -> DocxData:
    """Разобрать Word-файл (6.B.1).

    `.docx` — нативно через python-docx (абзацы с пометкой заголовков, таблицы).
    `.doc` (старый бинарный) — через внешний конвертер, только текст.
    """
    lower = (filename or "").lower()
    is_doc = lower.endswith(".doc") and not lower.endswith(".docx")
    if not (lower.endswith(".docx") or is_doc):
        raise DocumentError("Ожидается файл Word в формате .docx или .doc")
    _check_size(content, MAX_FILE_SIZE_WORD)

    if is_doc:
        return _parse_doc_legacy(content, filename)

    try:
        from docx import Document
    except ImportError as exc:  # pragma: no cover
        raise DocumentError("python-docx не установлен на сервере") from exc

    try:
        document = Document(io.BytesIO(content))
    except Exception as exc:  # noqa: BLE001 — python-docx кидает разнородные ошибки
        raise DocumentError(f"Не удалось прочитать файл как .docx: {exc}") from exc

    paragraphs: list[DocxParagraph] = []
    for para in document.paragraphs:
        text = (para.text or "").strip()
        if not text:
            continue
        style = (para.style.name if para.style else "") or "Normal"
        paragraphs.append(
            DocxParagraph(text=text, style=style, is_heading=style.lower().startswith("heading"))
        )

    tables: list[DocxTable] = []
    for table in document.tables:
        rows = [[(cell.text or "").strip() for cell in row.cells] for row in table.rows]
        tables.append(DocxTable(rows=rows))

    return DocxData(filename=filename, paragraphs=paragraphs, tables=tables)


@dataclass
class DocxDiffItem:
    op: str  # "added" | "removed" | "changed"
    old: str | None
    new: str | None


@dataclass
class DocxDiffResult:
    items: list[DocxDiffItem]
    truncated: bool

    @property
    def summary(self) -> dict[str, int]:
        out = {"added": 0, "removed": 0, "changed": 0}
        for it in self.items:
            out[it.op] += 1
        return out


def diff_docx(old: DocxData, new: DocxData) -> DocxDiffResult:
    """Сравнить два .docx по абзацам (6.B.3).

    Сравнение по тексту, форматирование не учитывается (осознанное упрощение
    первой итерации). Гранулярность — абзац: difflib по спискам абзацев.
    Соседние replace-блоки 1:1 трактуем как «изменён абзац».
    """
    old_paras = [p.text for p in old.paragraphs]
    new_paras = [p.text for p in new.paragraphs]

    matcher = difflib.SequenceMatcher(a=old_paras, b=new_paras, autojunk=False)
    items: list[DocxDiffItem] = []
    truncated = False

    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        if tag == "equal":
            continue
        if len(items) >= MAX_DIFF_ITEMS:
            truncated = True
            break
        if tag == "replace":
            old_block = old_paras[i1:i2]
            new_block = new_paras[j1:j2]
            # Сопоставимые 1:1 абзацы — «изменён»; излишек — удалён/добавлен.
            pairs = min(len(old_block), len(new_block))
            for k in range(pairs):
                items.append(DocxDiffItem(op="changed", old=old_block[k], new=new_block[k]))
            for text in old_block[pairs:]:
                items.append(DocxDiffItem(op="removed", old=text, new=None))
            for text in new_block[pairs:]:
                items.append(DocxDiffItem(op="added", old=None, new=text))
        elif tag == "delete":
            for text in old_paras[i1:i2]:
                items.append(DocxDiffItem(op="removed", old=text, new=None))
        elif tag == "insert":
            for text in new_paras[j1:j2]:
                items.append(DocxDiffItem(op="added", old=None, new=text))

    return DocxDiffResult(items=items[:MAX_DIFF_ITEMS], truncated=truncated)
