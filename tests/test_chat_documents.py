"""Тесты обработки файлов, приложенных к сообщению в чате (Этап 6)."""

from __future__ import annotations

import base64
import io

import pytest
from docx import Document
from openpyxl import Workbook, load_workbook

import app.state as state
from app.services import chat_documents as cd
from tests.conftest import ADMIN_EMAIL, ADMIN_PASSWORD, auth_headers


# ------------------------------ хелперы файлов ---------------------------


def _xlsx_bytes(headers: list[str], rows: list[list], *, sheet_name: str = "Sheet1") -> bytes:
    wb = Workbook()
    ws = wb.active
    ws.title = sheet_name
    ws.append(headers)
    for row in rows:
        ws.append(row)
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _docx_bytes(paragraphs: list[tuple[str, str]]) -> bytes:
    doc = Document()
    for text, style in paragraphs:
        doc.add_paragraph(text, style=style) if style else doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _xlsx_upload(name: str, content: bytes):
    return (name, content, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")


def _docx_upload(name: str, content: bytes):
    return (name, content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")


class _StubLLM:
    def summarize_text(self, text, **kw):
        return "КРАТКО: " + text[:20]


@pytest.fixture
def headers(client):
    return auth_headers(client, ADMIN_EMAIL, ADMIN_PASSWORD)


# =============== unit-тесты маршрутизации (handle_chat_files) =============


def test_route_single_xlsx_extracts():
    files = [("data.xlsx", _xlsx_bytes(["ID", "Имя"], [[1, "Иван"]]))]
    res = cd.handle_chat_files("", files, _StubLLM())
    assert "data.xlsx" in res.answer
    assert "Иван" in res.answer
    assert res.attachment is None


def test_route_single_xlsx_with_title_row():
    # Лист с объединённой «шапкой» в первой строке + настоящие заголовки ниже.
    wb = Workbook()
    ws = wb.active
    ws.title = "Лист 1 - Сотрудники"
    ws.append(["Сотрудники", None, None])  # титул (как объединённая ячейка)
    ws.append(["№", "ФИО", "Должность"])  # реальные заголовки
    ws.append([1, "Иванов И.И.", "Инженер"])
    ws.append([2, "Петров П.П.", "Техник"])
    buf = io.BytesIO()
    wb.save(buf)

    res = cd.handle_chat_files("покажи содержимое", [("emp.xlsx", buf.getvalue())], _StubLLM())
    # Должно увидеть 3 колонки, а не 1, и ФИО/должности попасть в таблицу.
    assert "3 колонок" in res.answer
    assert "ФИО" in res.answer and "Должность" in res.answer
    assert "Иванов И.И." in res.answer and "Инженер" in res.answer


def test_route_two_xlsx_default_diff():
    a = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"], [2, "Пётр"]])
    b = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"], [2, "Павел"]])
    res = cd.handle_chat_files("", [("a.xlsx", a), ("b.xlsx", b)], _StubLLM())
    assert "Различия между файлами" in res.answer
    assert "Пётр" in res.answer and "Павел" in res.answer


def test_route_merge_keyword_produces_file():
    a = _xlsx_bytes(["ID"], [[1]])
    b = _xlsx_bytes(["ID"], [[2]])
    res = cd.handle_chat_files("объедини эти файлы", [("a.xlsx", a), ("b.xlsx", b)], _StubLLM())
    assert res.attachment is not None
    assert res.attachment.filename == "merged.xlsx"
    # содержимое — валидный xlsx с двумя строками данных
    wb = load_workbook(io.BytesIO(base64.b64decode(res.attachment.content_base64)))
    assert wb.active.max_row == 3  # header + 2 rows


def test_route_merge_mismatch_no_file():
    a = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"]])
    b = _xlsx_bytes(["ID", "Фамилия"], [[2, "Петров"]])
    res = cd.handle_chat_files("объедини", [("a.xlsx", a), ("b.xlsx", b)], _StubLLM())
    assert res.attachment is None
    assert "структура" in res.answer.lower()


def test_route_single_docx_summary_uses_llm():
    f = [("doc.docx", _docx_bytes([("Большой текст документа.", "")]))]
    res = cd.handle_chat_files("сделай выжимку", f, _StubLLM())
    assert res.answer.startswith("**Выжимка")
    assert "КРАТКО:" in res.answer


def test_route_two_docx_diff():
    a = _docx_bytes([("Первый.", ""), ("Второй.", "")])
    b = _docx_bytes([("Первый.", ""), ("Второй изменён.", "")])
    res = cd.handle_chat_files("сравни", [("a.docx", a), ("b.docx", b)], _StubLLM())
    assert "Различия между файлами" in res.answer
    assert "«Второй.» → «Второй изменён.»" in res.answer


def test_docx_diff_attributes_lines_to_document():
    # Строка «000» есть только во втором документе — должна быть подписана им.
    one = _docx_bytes([("Общая строка.", ""), ("Ещё общая.", "")])
    two = _docx_bytes([("Общая строка.", ""), ("000", ""), ("Ещё общая.", "")])
    res = cd.handle_chat_files("сравни", [("1.docx", one), ("2.docx", two)], _StubLLM())
    assert "Только в документе «2.docx»" in res.answer
    assert "строка «000»" in res.answer
    # уникального в первом нет
    assert "Только в документе «1.docx»" not in res.answer


def test_route_mixed_types_rejected():
    with pytest.raises(cd.DocumentError):
        cd.handle_chat_files(
            "",
            [("a.xlsx", _xlsx_bytes(["ID"], [[1]])), ("b.docx", _docx_bytes([("x", "")]))],
            _StubLLM(),
        )


def test_route_unsupported_extension_rejected():
    with pytest.raises(cd.DocumentError):
        cd.handle_chat_files("", [("notes.txt", b"hello")], _StubLLM())


# ===================== старый бинарный .doc (Этап 6, #7) =================
#
# Настоящий .doc в тестах не сгенерировать без Word, поэтому конвертер
# `_doc_to_text` подменяем — проверяем интеграцию (маршрутизация .doc как
# Word, построение абзацев, обработка ошибок), а не сам внешний бинарь.

_DOC_TEXT = "Заголовок отчёта\nПервый абзац текста.\nВторой абзац текста."


def test_doc_routed_as_word_summary(monkeypatch):
    monkeypatch.setattr(
        "app.services.document_processing._doc_to_text", lambda content: _DOC_TEXT
    )
    res = cd.handle_chat_files("сделай выжимку", [("old.doc", b"\xd0\xcf\x11\xe0fake")], _StubLLM())
    assert res.answer.startswith("**Выжимка")
    assert "КРАТКО:" in res.answer


def test_doc_routed_as_word_extract(monkeypatch):
    monkeypatch.setattr(
        "app.services.document_processing._doc_to_text", lambda content: _DOC_TEXT
    )
    res = cd.handle_chat_files("извлеки данные", [("old.doc", b"fake")], _StubLLM())
    assert "Первый абзац" in res.answer


def test_doc_and_docx_compared_together(monkeypatch):
    # .doc и .docx — оба Word, должны группироваться и сравниваться.
    monkeypatch.setattr(
        "app.services.document_processing._doc_to_text",
        lambda content: "Первый.\nВторой.",
    )
    new = _docx_bytes([("Первый.", ""), ("Второй изменён.", "")])
    res = cd.handle_chat_files("сравни", [("old.doc", b"fake"), ("new.docx", new)], _StubLLM())
    assert "Различия между файлами" in res.answer
    assert "«Второй.» → «Второй изменён.»" in res.answer


def test_doc_empty_text_raises(monkeypatch):
    monkeypatch.setattr(
        "app.services.document_processing._doc_to_text", lambda content: "   \n  \n"
    )
    with pytest.raises(cd.DocumentError, match="не найдено текста"):
        cd.handle_chat_files("выжимка", [("old.doc", b"fake")], _StubLLM())


def test_doc_no_converter_available(monkeypatch):
    # Ни antiword, ни textutil не найдены — понятная ошибка с просьбой пересохранить.
    monkeypatch.setattr("app.services.document_processing.shutil.which", lambda name: None)
    with pytest.raises(cd.DocumentError, match="конвертера|docx"):
        cd.handle_chat_files("выжимка", [("old.doc", b"\xd0\xcf\x11\xe0fake")], _StubLLM())


def test_doc_to_text_uses_antiword(monkeypatch):
    import subprocess

    from app.services import document_processing as dp

    calls = {}

    def fake_which(name):
        return "/usr/bin/antiword" if name == "antiword" else None

    def fake_run(cmd, **kw):
        calls["cmd"] = cmd
        return subprocess.CompletedProcess(cmd, 0, stdout="Текст из antiword".encode(), stderr=b"")

    monkeypatch.setattr(dp.shutil, "which", fake_which)
    monkeypatch.setattr(dp.subprocess, "run", fake_run)
    assert dp._doc_to_text(b"fake") == "Текст из antiword"
    assert calls["cmd"][0] == "/usr/bin/antiword"


# ===================== раздельные лимиты размера (Этап 6, #2) =============


def test_size_limits_differ_by_type():
    from app.services import document_processing as dp

    assert dp.MAX_FILE_SIZE_XLSX < dp.MAX_FILE_SIZE_WORD  # Excel строже Word


def test_oversize_xlsx_rejected(monkeypatch):
    from app.services import document_processing as dp

    monkeypatch.setattr(dp, "MAX_FILE_SIZE_XLSX", 100)  # 100 байт
    big = _xlsx_bytes(["ID"], [[1]])  # заведомо больше 100 байт
    with pytest.raises(cd.DocumentError, match="слишком большой"):
        cd.handle_chat_files("что тут", [("big.xlsx", big)], _StubLLM())


# ============ объединение Excel по именам колонок, #4 (порядок свободный) =


def test_merge_reorders_columns_by_name():
    # Одинаковый набор колонок, разный порядок — должно объединиться с переносом
    # строк к порядку первого файла.
    a = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"]])
    b = _xlsx_bytes(["Имя", "ID"], [["Пётр", 2]])  # колонки переставлены
    res = cd.handle_chat_files("объедини", [("a.xlsx", a), ("b.xlsx", b)], _StubLLM())
    assert res.attachment is not None, res.answer
    wb = load_workbook(io.BytesIO(base64.b64decode(res.attachment.content_base64)))
    ws = wb.active
    values = [[c.value for c in row] for row in ws.iter_rows()]
    assert values[0] == ["ID", "Имя"]  # порядок первого файла
    assert ["2", "Пётр"] in values  # строка b переставлена под этот порядок


def test_merge_case_insensitive_headers():
    a = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"]])
    b = _xlsx_bytes([" id ", "ИМЯ"], [[2, "Пётр"]])  # регистр/пробелы
    res = cd.handle_chat_files("объедини", [("a.xlsx", a), ("b.xlsx", b)], _StubLLM())
    assert res.attachment is not None, res.answer


def test_merge_missing_column_reported():
    a = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"]])
    b = _xlsx_bytes(["ID", "Фамилия"], [[2, "Петров"]])
    res = cd.handle_chat_files("объедини", [("a.xlsx", a), ("b.xlsx", b)], _StubLLM())
    assert res.attachment is None
    assert "не хватает колонок" in res.answer or "лишние колонки" in res.answer


# ===================== пропорциональная длина выжимки (Этап 6, #6) ========


# ============ сравнение бланков-форм, а не таблиц (Этап 6, 6.A.3) =========


def _form_xlsx(pairs: list[tuple[str, str, str]]) -> bytes:
    """Бланк-форма: разрежённый лист, пары «подпись (B) — значение (F)» по строкам.

    pairs: список (row_label, value, _). Подпись кладём в колонку B, значение — в F,
    строки с пропусками — чтобы плотность была низкой (как у реальной заявки).
    """
    wb = Workbook()
    ws = wb.active
    row = 3
    for label, value, _ in pairs:
        ws.cell(row=row, column=2, value=label)  # B — подпись
        ws.cell(row=row, column=6, value=value)  # F — значение
        row += 2  # пропуск строки → разрежённость
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_form_diff_uses_labels_not_positions():
    a = _form_xlsx([("Ф.И.О. работника", "Петров Иван Иванович", ""), ("Подразделение", "ИТ", "")])
    b = _form_xlsx(
        [("Ф.И.О. работника", "Петров Иван Иванович, 89123451234", ""), ("Подразделение", "ИТ", "")]
    )
    res = cd.handle_chat_files("сравни", [("Заявка1.xlsx", a), ("Заявка2.xlsx", b)], _StubLLM())
    assert "Различия между бланками" in res.answer
    assert "Ф.И.О. работника" in res.answer  # подпись, а не «строка N» / значение-колонка
    assert "89123451234" in res.answer
    assert "изменено полей — 1" in res.answer


def test_form_diff_identical():
    a = _form_xlsx([("Поле", "Значение", "")])
    res = cd.handle_chat_files("сравни", [("a.xlsx", a), ("b.xlsx", a)], _StubLLM())
    assert "идентичны" in res.answer


def test_dense_table_still_uses_table_diff():
    # Плотная таблица-реестр не должна попасть в режим форм.
    a = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"], [2, "Пётр"]])
    b = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"], [2, "Павел"]])
    res = cd.handle_chat_files("сравни", [("a.xlsx", a), ("b.xlsx", b)], _StubLLM())
    assert "Различия между бланками" not in res.answer  # табличный режим
    assert "изменено ячеек" in res.answer


def test_table_diff_attributes_rows_to_document():
    # У b есть лишняя строка (ID 3) — должна быть подписана документом b.
    a = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"], [2, "Пётр"]])
    b = _xlsx_bytes(["ID", "Имя"], [[1, "Иван"], [2, "Пётр"], [3, "Семён"]])
    res = cd.handle_chat_files("сравни", [("реестр1.xlsx", a), ("реестр2.xlsx", b)], _StubLLM())
    assert "Только в документе «реестр2.xlsx»" in res.answer
    assert "Семён" in res.answer
    assert "Только в документе «реестр1.xlsx»" not in res.answer


def test_summary_length_instruction_scales():
    from app.llm import LLMService

    short = LLMService._summary_length_instruction("слово " * 50)
    medium = LLMService._summary_length_instruction("слово " * 1000)
    long_ = LLMService._summary_length_instruction("слово " * 5000)
    huge = LLMService._summary_length_instruction("слово " * 20000)
    assert "2-3 предложения" in short
    assert "100-150" in medium
    assert "200-300" in long_
    assert "300-400" in huge


# ===================== интеграция через /api/ask-files ===================


def test_endpoint_requires_auth(client):
    resp = client.post(
        "/api/ask-files",
        data={"message": "что тут"},
        files={"files": _xlsx_upload("x.xlsx", _xlsx_bytes(["a"], [[1]]))},
    )
    assert resp.status_code == 401


def test_endpoint_extract_without_conversation(client, headers):
    resp = client.post(
        "/api/ask-files",
        data={"message": "что в файле"},
        files={"files": _xlsx_upload("data.xlsx", _xlsx_bytes(["ID", "Имя"], [[1, "Иван"]]))},
        headers=headers,
    )
    assert resp.status_code == 200, resp.text
    data = resp.json()
    assert "Иван" in data["answer"]
    assert data["attachment"] is None


def test_endpoint_summary_persists_in_conversation(client, headers, monkeypatch):
    monkeypatch.setattr(state.llm_service, "summarize_text", lambda text, **kw: "РЕЗЮМЕ ТУТ")
    conv = client.post("/api/conversations", json={}, headers=headers).json()
    resp = client.post(
        "/api/ask-files",
        data={"message": "выжимка", "conversation_id": str(conv["id"])},
        files={"files": _docx_upload("doc.docx", _docx_bytes([("Текст документа.", "")]))},
        headers=headers,
    )
    assert resp.status_code == 200, resp.text
    assert "РЕЗЮМЕ ТУТ" in resp.json()["answer"]

    # сообщения (вопрос с пометкой файла + ответ) сохранены в беседе
    detail = client.get(f"/api/conversations/{conv['id']}", headers=headers).json()
    roles = [m["role"] for m in detail["messages"]]
    assert roles == ["user", "assistant"]
    assert "📎 doc.docx" in detail["messages"][0]["content"]
